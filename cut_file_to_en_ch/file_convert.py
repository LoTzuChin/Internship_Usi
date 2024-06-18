import docx
import re

def process_paragraphs(doc, pattern, pattern_char, file):
    """
    將 docx內的文字複寫進 txt
    :param doc: 輸入的 docx檔
    :param pattern: 用於匹配數字編號的正規表示式
    :param pattern_char: 用於匹配字符編號的正規表達式
    :param file: 被寫入的文件
    """
    indent = 0  # 縮排數量
    temp = 0  # 縮排數量的臨時記憶空間

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        # 抓取符合 pattern 的文字内容
        match = pattern.match(paragraph.text)
        # 抓取符合 pattern_char 的文字内容
        match_char = pattern_char.match(paragraph.text)

        # 根據抓取的內容判斷要縮進幾排
        # 優先判斷符合 pattren的
        if match:
            numbering = match.group(1)  # 获取类似于 '3.1' 的编号

            # 根據 numbering得形式，決定縮進的數量
            blocks = [block.strip() for block in numbering.split('.') if block.strip()]
            indent = len(blocks) - 1

            # 暫存縮進數量
            temp = indent

        elif match_char:
            # 如果符合的是 pattern_char，則將在上一層的縮排數量下+1
            # 原因：符合 pattern_char的都是依附在 pattern下做補充的
            indent = temp + 1

        file.write(indent * '\t' + paragraph.text + '\n\n\n')

def process_table(table, index, file):
    """
    將 docx內的表格轉為table llama的形式
    :param table: 輸入之 docx內的表格
    :param index: 紀錄表格編號
    :param file: 被寫入的文件
    :return: 返回表格編號，讓編號得以延續
    """
    count = 1
    # [TLE] 紀錄的是表格編號，其中數字用兩位數表示，不足則以0填滿
    file.write(f'[TLE] Table {str(index).zfill(2)}')

    # 遍歷表格，完成改寫
    # ex: [TLE] Table 01 [TAB] | Table name | Table name | [SEP] | field name | field name | [SEP] | ... | ... |
    for row in table.rows:
        if count == 1:
            # [TAB] 用於表是表格名稱
            file.write(' [TAB] ')
            count += 1
        else:
            # [SEP] 表其他欄位，包含欄位名稱及其內容
            file.write(' [SEP] ')

        for cell in row.cells:
            cell_text = cell.text.replace('\n', '')
            file.write('| ' + cell_text + ' ')
        file.write('|')
    file.write('\n\n\n')

    # 下一個表格的編號
    index += 1

    return index


if __name__ == "__main__":

    # 輸入之docx
    doc = docx.Document("AT-0801-E217-AR.docx")

    index = 1  # 表格编号

    # 因為要通過編號進行計算需要幾個tab，所以先進行編號的正規化
    # pattern 正規化的部分是單純數字的編號，如：1., 1.1, 1.1.1
    # pattern_char 正規化部分是小編號部分，相比上方呈現形式較不確定，如：(1), (a), a.
    pattern = re.compile(r'^(\d\.|\d+(\.\d+)*)\s')
    pattern_char = re.compile(r'^(\(\d\)|(\w\.)|(\(\w)\))\s')

    # 將 docx內容複寫進 txt檔
    with open('output.txt', 'w', encoding='utf-8') as f:

        # 文字內容
        process_paragraphs(doc, pattern, pattern_char, f)

        # 表格內容
        for table in doc.tables:
            index = process_table(table, index, f)
