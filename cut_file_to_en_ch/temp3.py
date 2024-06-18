import docx
import re

def contains_chinese(text):
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            return True
    return False

def process_paragraphs(pattern, pattern_char, ch_file, en_file, indent, temp):
    """
    將 docx內的文字複寫進 txt
    :param pattern: 用於匹配數字編號的正規表示式
    :param pattern_char: 用於匹配字符編號的正規表達式
    """

    # 抓取符合 pattern 的文字内容
    match = pattern.match(paragraph.text)
    # 抓取符合 pattern_char 的文字内容
    match_char = pattern_char.match(paragraph.text)


    # 根據抓取的內容判斷要縮進幾排
    # 優先判斷符合 pattren的
    if match:
        numbering = match.group(1)  # 获取类似于 '3.1' 的编号
        number = str(numbering)
        # 根據 numbering得形式，決定縮進的數量
        blocks = [block.strip() for block in numbering.split('.') if block.strip()]
        indent = len(blocks) - 1

        # 暫存縮進數量
        temp = indent

    elif match_char:
        numbering = match_char.group(1)
        number = str(numbering)
        # print(number)
        # print(type(numbering))
        # 如果符合的是 pattern_char，則將在上一層的縮排數量下+1
        # 原因：符合 pattern_char的都是依附在 pattern下做補充的
        indent = temp + 1

    else:
        numbering = ''
        number = str(numbering)

    if contains_chinese(paragraph.text):
        # 如果包含中文字符，写入中文文件
        ch_file.write(indent * '\t' + paragraph.text + '\n\n\n')
    else:
        # print(numbering)
        # 否则，写入英文文件
        en_file.write(indent * '\t' + number + paragraph.text + '\n\n\n')

    return indent, temp

def process_table(table, index, file):
    """
    將 docx內的表格轉為table llama的形式
    :param table: 輸入之 docx內的表格
    :param index: 紀錄表格編號
    :param file: 被寫入的文件
    :return: 返回表格編號，讓編號得以延續
    """
    count = 1
    # [TLE] 紀錄的是表格編號，其中數字用兩位數表示，不足則以0填滿
    file.write(f'[TLE] Table {str(index).zfill(2)}')

    # 遍歷表格，完成改寫
    # ex: [TLE] Table 01 [TAB] | Table name | Table name | [SEP] | field name | field name | [SEP] | ... | ... |
    for row in table.rows:
        if count == 1:
            # [TAB] 用於表是表格名稱
            file.write(' [TAB] ')
            count += 1
        else:
            # [SEP] 表其他欄位，包含欄位名稱及其內容
            file.write(' [SEP] ')

        for cell in row.cells:
            cell_text = cell.text.replace('\n', '')
            file.write('| ' + cell_text + ' ')
        file.write('|')
    file.write('\n\n\n')

    # 下一個表格的編號
    index += 1

    return index


if __name__ == "__main__":

    # 輸入之docx
    doc = docx.Document("AT-0801-E217-AR.docx")

    index = 1  # 表格编号
    indent = 0  # 縮排數量
    temp = 0  # 縮排數量的臨時記憶空間

    # 因為要通過編號進行計算需要幾個tab，所以先進行編號的正規化
    # pattern 正規化的部分是單純數字的編號，如：1., 1.1, 1.1.1
    # pattern_char 正規化部分是小編號部分，相比上方呈現形式較不確定，如：(1), (a), a.
    pattern = re.compile(r'^(\d\.|\d+(\.\d+)*)\s')
    pattern_char = re.compile(r'^(\(\d\)|(\w\.)|(\(\w)\))\s')

    # 將 docx內容複寫進 txt檔
    with open('chinese.txt', 'w', encoding='utf-8') as ch_f, open('english.txt', 'w', encoding='utf-8') as en_f:

        # 文字內容
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
            indent, temp = process_paragraphs(pattern, pattern_char, ch_f, en_f, indent, temp)

        # 表格內容
        # for table in doc.tables:
        #     index = process_table(table, index, f)
